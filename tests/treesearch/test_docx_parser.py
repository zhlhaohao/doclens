"""tests for treesearch.parsers.docx_parser.

回归 docx 段落被合并成单个 markdown 段落的 bug：
node.text 必须用 ``\n\n`` 分段，preview synthesizer 才能渲染成独立段落。
"""
import asyncio
from pathlib import Path

from treesearch.parsers.docx_parser import docx_to_tree


def _run(coro):
    return asyncio.run(coro)


def _write_docx_with_heading_and_paragraphs(path: Path) -> None:
    """生成一份「1 个 Heading 1 + 3 个 Normal 段落（夹 1 个空段）+ 1 个表格」的 docx。

    模拟联通应急预案结构：Word heading 样式 + 正文段落 + 空段 + 表格混排。
    """
    from docx import Document

    doc = Document()
    doc.add_heading("场景一：示例场景", level=1)
    doc.add_paragraph("第一段正文内容，描述判断条件。")
    doc.add_paragraph("")  # 空段（docx 常用作视觉间距）
    doc.add_paragraph("第二段正文内容，描述业务影响。")
    doc.add_paragraph("第三段正文内容，描述处置步骤。")

    # 加一个简单表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列A"
    table.cell(0, 1).text = "列B"
    table.cell(1, 0).text = "v1"
    table.cell(1, 1).text = "v2"

    doc.save(str(path))


def test_docx_paragraphs_are_double_newline_separated(tmp_path: Path):
    """Heading 路径：node.text 必须用 ``\n\n`` 分隔每个非空段落。

    回归：原实现 ``"\\n".join(lines[start:end])`` 导致 preview 渲染时
    所有段落被合并成一个 markdown 段落。
    """
    docx_path = tmp_path / "sample.docx"
    _write_docx_with_heading_and_paragraphs(docx_path)

    result = _run(
        docx_to_tree(
            str(docx_path),
            if_add_node_summary=False,
            if_add_node_text=True,
        )
    )

    structure = result["structure"]
    assert len(structure) == 1, "应当只有一个顶级 Heading 1 节点"
    node = structure[0]
    text = node.get("text", "")

    # 用 \n\n 切分，应当得到：heading 文本 + 3 段正文 + 1 个表格块 = 5 段
    paragraphs = text.split("\n\n")
    assert len(paragraphs) == 5, (
        f"期望 5 段（heading + 3 正文 + 表格），实际 {len(paragraphs)} 段；"
        f"text repr={text!r}"
    )

    # heading 段（首段）
    assert paragraphs[0] == "场景一：示例场景"
    # 3 段正文，顺序保留
    assert paragraphs[1] == "第一段正文内容，描述判断条件。"
    assert paragraphs[2] == "第二段正文内容，描述业务影响。"
    assert paragraphs[3] == "第三段正文内容，描述处置步骤。"
    # 表格段：表头 + 数据行，用单 \n 分隔
    table_block = paragraphs[4]
    table_lines = table_block.split("\n")
    assert len(table_lines) == 3, (
        f"表格应为 3 行（header + separator + data），实际 {len(table_lines)} 行；"
        f"table_block={table_block!r}"
    )
    # GFM md 表格：header 行 + separator 行 + 数据行
    assert "列A" in table_lines[0] and "列B" in table_lines[0]
    assert table_lines[1] == "| --- | --- |", f"separator 行应为 '| --- | --- |'，实际 {table_lines[1]!r}"
    assert "v1" in table_lines[2] and "v2" in table_lines[2]


def _write_docx_without_heading_styles(path: Path) -> None:
    """生成一份**没有任何 Heading 样式**的 docx，仅 3 个 Normal 段落。

    用于触发 docx_to_tree 的 text_to_tree fallback 路径。
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段普通正文。")
    doc.add_paragraph("第二段普通正文。")
    doc.add_paragraph("第三段普通正文。")
    doc.save(str(path))


def test_docx_fallback_path_also_separates_paragraphs(tmp_path: Path):
    """Fallback 路径：无 Word heading 样式时，text_to_tree 收到的 text_content
    也必须用 ``\\n\\n`` 分段，否则 fallback 场景下段落同样会合并。"""
    docx_path = tmp_path / "plain.docx"
    _write_docx_without_heading_styles(docx_path)

    result = _run(
        docx_to_tree(
            str(docx_path),
            if_add_node_summary=False,
            if_add_node_text=True,
        )
    )

    structure = result["structure"]
    # fallback 路径会通过 _detect_headings 找到子标题或合成单根节点；
    # 我们只关心：在生成的所有节点 text 里，至少存在一个 \n\n。
    all_text = []
    for n in structure:
        all_text.append(n.get("text", ""))
        for c in n.get("nodes") or []:
            all_text.append(c.get("text", ""))

    joined = "\n\n".join(t for t in all_text if t)
    assert "\n\n" in joined, (
        f"fallback 路径未生成 \\n\\n 分段，joined repr={joined!r}"
    )


# ============================================================================
# docx 表格 → GFM markdown table 渲染
# ============================================================================


def _write_docx_with_table_only(path: Path, header_cells, data_rows) -> None:
    """生成一份只含一张表格、无任何 Heading 的 docx（便于直接测 _table_to_text）。"""
    from docx import Document

    doc = Document()
    ncols = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    for i, h in enumerate(header_cells):
        table.cell(0, i).text = h
    for r, row_data in enumerate(data_rows, start=1):
        for i, v in enumerate(row_data):
            table.cell(r, i).text = v
    doc.save(str(path))


def test_table_to_text_outputs_github_flavored_md_table(tmp_path: Path):
    """`_table_to_text` 必须输出 GFM md table 语法。

    回归：原实现只输出 `' | '.join(cells)` 多行，缺第二行 `| --- |` separator，
    导致前端 `<md-viewer>` 不识别为表格。
    """
    from treesearch.parsers.docx_parser import _table_to_text

    docx_path = tmp_path / "tbl.docx"
    _write_docx_with_table_only(
        docx_path,
        header_cells=["操作步骤", "风险", "处置措施"],
        data_rows=[
            ["删除Pod重建", "服务不可用", "确保多副本"],
            ["修改镜像标签", "版本异常", "测试环境验证"],
        ],
    )

    # 用 python-docx 打开，取出第一张 table
    from docx import Document
    doc = Document(str(docx_path))
    table = doc.tables[0]

    output = _table_to_text(table)
    lines = output.split("\n")

    # 期望：1 header + 1 separator + 2 data = 4 行
    assert len(lines) == 4, f"期望 4 行，实际 {len(lines)}：{output!r}"
    # header 行规范：以 | 开头、结尾，单元格 | 分隔
    assert lines[0] == "| 操作步骤 | 风险 | 处置措施 |", f"header 行不符：{lines[0]!r}"
    # separator 行
    assert lines[1] == "| --- | --- | --- |", f"separator 行不符：{lines[1]!r}"
    # 数据行
    assert lines[2] == "| 删除Pod重建 | 服务不可用 | 确保多副本 |", f"data row 1 不符：{lines[2]!r}"
    assert lines[3] == "| 修改镜像标签 | 版本异常 | 测试环境验证 |", f"data row 2 不符：{lines[3]!r}"


def test_table_to_text_escapes_pipe_in_cell(tmp_path: Path):
    """单元格里的 `|` 必须转义为 `\\|`，否则会破坏 md 表格结构。"""
    from treesearch.parsers.docx_parser import _table_to_text

    docx_path = tmp_path / "pipe.docx"
    _write_docx_with_table_only(
        docx_path,
        header_cells=["h1", "h2"],
        data_rows=[["a|b", "c|d|e"]],
    )

    from docx import Document
    table = Document(str(docx_path)).tables[0]

    output = _table_to_text(table)
    # 数据行：`|` 都被转义为 `\|`
    data_line = output.split("\n")[2]
    assert data_line == "| a\\|b | c\\|d\\|e |", f"data 行 | 转义失败：{data_line!r}"


def test_table_to_text_collapses_newline_in_cell(tmp_path: Path):
    """单元格内的换行（docx cell 多段落）替换为空格，防 md 表格行被截断。"""
    from treesearch.parsers.docx_parser import _table_to_text

    docx_path = tmp_path / "nl.docx"
    _write_docx_with_table_only(
        docx_path,
        header_cells=["h1", "h2"],
        data_rows=[["line1", "normal"]],
    )
    # 给 (0,0) cell 追加一段，制造 cell 内换行
    from docx import Document
    doc = Document(str(docx_path))
    cell = doc.tables[0].cell(1, 0)
    cell.add_paragraph("line2")
    doc.save(str(docx_path))

    # 重新打开（避免 stale 对象）
    table = Document(str(docx_path)).tables[0]
    output = _table_to_text(table)
    data_line = output.split("\n")[2]
    # cell 内换行 → 空格
    assert data_line == "| line1 line2 | normal |", f"换行替换失败：{data_line!r}"


def test_docx_table_renders_as_md_table_in_preview(tmp_path: Path):
    """端到端：docx_to_tree → render_tree_to_md，preview md 必须含 `| --- |` 分隔行。"""
    docx_path = tmp_path / "doc_with_tbl.docx"
    _write_docx_with_table_only(
        docx_path,
        header_cells=["h1", "h2"],
        data_rows=[["v1", "v2"]],
    )
    # 加一个 Heading 1 让 docx_to_tree 走 heading 路径
    from docx import Document
    doc = Document(str(docx_path))
    # 在表格前插入一个 Heading 1（python-docx 不直接支持 reordering，重新构造）
    new_doc = Document()
    new_doc.add_heading("标题", level=1)
    new_doc.add_paragraph("正文一段。")
    # 复制原表格过来
    tbl_xml = doc.tables[0]._element
    new_doc._body._element.append(tbl_xml)
    new_doc.save(str(docx_path))

    result = _run(
        docx_to_tree(
            str(docx_path),
            if_add_node_summary=False,
            if_add_node_text=True,
        )
    )
    from doclens.web_v2.preview_synthesizer import render_tree_to_md
    md = render_tree_to_md(result["structure"], source_type="docx")

    # md 应包含 GFM separator 行
    assert "| --- | --- |" in md, f"preview md 未包含表格 separator 行，md={md!r}"
    # 数据行也在
    assert "| v1 | v2 |" in md, f"preview md 未包含表格数据行，md={md!r}"
