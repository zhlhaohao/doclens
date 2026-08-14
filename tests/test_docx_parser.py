# -*- coding: utf-8 -*-
"""docx_parser 列表符号保留测试。

背景：``para.text`` 不含 Word 渲染期生成的项目符号/编号，列表曾退化为普通
段落。``_list_marker`` 检测段落级 ``w:numPr``（Word 常规编码）或 List 系
样式名，补回 ``- `` 前缀（含 ilvl 缩进），与 anydoc 输出口径对齐。
"""
import pytest

docx = pytest.importorskip("docx", reason="python-docx not installed")

from docx.oxml.ns import qn  # noqa: E402
from lxml import etree  # noqa: E402

from treesearch.parsers.docx_parser import _extract_docx_headings  # noqa: E402


def _add_numpr(para, ilvl: int = 0, num_id: int = 2):
    """给段落注入段落级 w:numPr（模拟 Word 列表的真实编码）。"""
    pPr = para._p.get_or_add_pPr()
    numPr = etree.SubElement(pPr, qn("w:numPr"))
    ilvl_el = etree.SubElement(numPr, qn("w:ilvl"))
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = etree.SubElement(numPr, qn("w:numId"))
    numId_el.set(qn("w:val"), str(num_id))


@pytest.fixture()
def list_docx(tmp_path):
    """构造含标题、普通段落、两级列表的 docx。"""
    doc = docx.Document()
    doc.add_heading("章节一", level=1)
    doc.add_paragraph("普通段落")
    p1 = doc.add_paragraph("一级条目")
    _add_numpr(p1, ilvl=0)
    p2 = doc.add_paragraph("二级条目")
    _add_numpr(p2, ilvl=1)
    doc.add_paragraph("样式列表项", style="List Bullet")
    fp = tmp_path / "list.docx"
    doc.save(str(fp))
    return str(fp)


def test_list_paragraphs_get_bullet_prefix(list_docx):
    headings, lines = _extract_docx_headings(list_docx)
    assert [h["title"] for h in headings] == ["章节一"]  # 标题不被加前缀
    assert "普通段落" in lines
    assert "- 一级条目" in lines
    assert "  - 二级条目" in lines  # ilvl=1 → 两空格缩进
    assert "- 样式列表项" in lines  # List Bullet 样式（无段落级 numPr）也识别


def test_non_list_paragraphs_unchanged(list_docx):
    _, lines = _extract_docx_headings(list_docx)
    assert lines[0] == "章节一"
    assert lines[1] == "普通段落"


async def test_list_survives_tree_build(list_docx):
    """e2e：列表符号进入树的节点文本，可被搜索/预览消费。"""
    from treesearch.parsers.docx_parser import docx_to_tree

    result = await docx_to_tree(list_docx, if_add_node_text=True)
    node = result["structure"][0]
    assert "- 一级条目" in node["text"]
    assert "  - 二级条目" in node["text"]
